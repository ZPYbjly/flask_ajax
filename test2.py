from docx import Document
from docx.shared import Inches

# 创建一个新的 Word 文档
document = Document()

# 添加标题
document.add_heading('Document Title', 0)

# 添加段落
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# 添加一个表格
table = document.add_table(rows=3, cols=3)
for row in table.rows:
    for cell in row.cells:
        cell.text = 'Cell'

# 添加一张图片
# document.add_picture('picture.png', width=Inches(1.25))

# 保存文档
document.save('document.docx')