#  rangetype   province  city    district   typeselect  csvselect
# 创建postgresql

# create

# 基本引用 通过 Document 可以创建一个文档对象

from docx import Document
# 从文件创建文档对象
document = Document('C:/Users/pengyu/Desktop/flask_ajax/template.docx')
# C:/Users/pengyu/Desktop/flask_ajax/template.docx
# 显示每段的内容
for p in document.paragraphs:
    print(p.text)
# 添加段落
document.add_paragraph('这是新的段落内容')
# 保存文档
document.save('C:/Users/pengyu/Desktop/flask_ajax/template2.docx')